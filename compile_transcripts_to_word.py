from docx import Document
import os
import re

def format_text(text):
    """Remove unnecessary line breaks, adjust capitalization, and ensure single spaces."""
    # Join lines into a single block of text
    text = ' '.join(text.split('\n'))
    # Adjust capitalization after full stops
    text = re.sub(r'\. ([a-z])', lambda x: '. ' + x.group(1).upper(), text)
    # Replace multiple spaces with a single space
    text = re.sub(r'\s+', ' ', text)
    return text

def compile_transcripts_to_word(transcripts_folder, output_file):
    document = Document()
    
    for filename in sorted(os.listdir(transcripts_folder), key=lambda x: int(x.split('_')[0])):
        if filename.endswith('.txt'):
            # Extract playlist title from filename, assuming format "transcript_{Playlist Title}.txt"
            playlist_title = filename.split('_', 2)[-1].replace('transcript_', '').replace('.txt', '')
            document.add_heading(playlist_title, level=1)
            
            with open(os.path.join(transcripts_folder, filename), 'r', encoding='utf-8') as file:
                content = file.read()
                # Split content by asterisks lines, assuming each title is surrounded by asterisks
                parts = re.split(r'\n\*(.*?)\*\n', content)
                for i in range(1, len(parts), 2):
                    video_title = parts[i]
                    video_transcript = parts[i+1]
                    document.add_heading(video_title, level=2)
                    formatted_text = format_text(video_transcript)
                    document.add_paragraph(formatted_text)
    
    document.save(output_file)
    print(f"Compiled Word document saved as: {output_file}")

if __name__ == "__main__":
    transcripts_folder = "txt_transcripts"  # Folder containing the .txt files
    output_file = "Playlist_Transcripts.docx"  # Output Word file name
    compile_transcripts_to_word(transcripts_folder, output_file)
